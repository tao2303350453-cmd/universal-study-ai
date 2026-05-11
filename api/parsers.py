"""
File parsing module for Review Assistant.
Supports: PDF, DOCX, XLSX, XLS, CSV, TXT
"""

import io
import logging

logger = logging.getLogger(__name__)


def parse_file(content: bytes, filename: str) -> str:
    """
    Parse uploaded file bytes and return extracted text.
    
    Args:
        content: Raw file bytes
        filename: Original filename (used to determine type)
    
    Returns:
        Extracted text string
    
    Raises:
        ValueError: If file type is unsupported or parsing fails
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    try:
        if ext == "txt":
            return _parse_txt(content)
        elif ext == "pdf":
            return _parse_pdf(content)
        elif ext == "docx":
            return _parse_docx(content)
        elif ext in ("xlsx", "xls"):
            return _parse_excel(content)
        elif ext == "csv":
            return _parse_csv(content)
        else:
            raise ValueError(f"不支持的文件格式：.{ext}，请上传 PDF、Word、Excel、CSV 或 TXT 文件。")
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Error parsing {filename}: {e}")
        raise ValueError(f"文件解析失败：{str(e)}")


# ─────────────────────────── Individual parsers ───────────────────────────

def _parse_txt(content: bytes) -> str:
    """Parse plain text file, trying multiple encodings."""
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    """Parse PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ValueError("PyMuPDF 未安装，请运行：pip install pymupdf")

    doc = fitz.open(stream=content, filetype="pdf")
    pages_text = []

    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text.strip():
            pages_text.append(f"[第 {page_num} 页]\n{text.strip()}")

    doc.close()

    if not pages_text:
        raise ValueError("PDF 文件为空或无法提取文本（可能是扫描版，暂不支持 OCR）。")

    return "\n\n".join(pages_text)


def _parse_docx(content: bytes) -> str:
    """Parse Word DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx 未安装，请运行：pip install python-docx")

    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading style info
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
            else:
                parts.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    if not parts:
        raise ValueError("Word 文档内容为空。")

    return "\n\n".join(parts)


def _parse_excel(content: bytes) -> str:
    """Parse Excel file using pandas."""
    try:
        import pandas as pd
    except ImportError:
        raise ValueError("pandas 未安装，请运行：pip install pandas openpyxl")

    excel_file = pd.ExcelFile(io.BytesIO(content))
    parts = []

    for sheet_name in excel_file.sheet_names:
        df = excel_file.parse(sheet_name)
        df = df.dropna(how="all").dropna(axis=1, how="all")

        if df.empty:
            continue

        parts.append(f"## 工作表：{sheet_name}")
        parts.append(df.to_string(index=False, max_rows=500))

    if not parts:
        raise ValueError("Excel 文件内容为空。")

    return "\n\n".join(parts)


def _parse_csv(content: bytes) -> str:
    """Parse CSV file using pandas with encoding detection."""
    try:
        import pandas as pd
    except ImportError:
        raise ValueError("pandas 未安装，请运行：pip install pandas")

    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312", "latin-1"):
        try:
            df = pd.read_csv(io.BytesIO(content), encoding=encoding)
            df = df.dropna(how="all").dropna(axis=1, how="all")
            if df.empty:
                raise ValueError("CSV 文件内容为空。")
            return df.to_string(index=False, max_rows=1000)
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue

    raise ValueError("无法解析 CSV 文件，请确认编码格式。")
